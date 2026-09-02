"""Text extraction from resume files.

Uses `pypdf` (PyPDF2's maintained successor — the old code imported PyPDF2,
which is deprecated and not installed) and `python-docx`. Legacy `.doc` and
`.rtf` files are detected and reported rather than silently mis-parsed.
"""
from __future__ import annotations

import re
import unicodedata
from dataclasses import dataclass, field
from pathlib import Path

SUPPORTED_SUFFIXES = {".pdf", ".docx", ".txt"}
UNSUPPORTED_LEGACY = {".doc", ".rtf", ".odt", ".pages"}


class ExtractionError(RuntimeError):
    """Raised when a file cannot be turned into usable text."""


@dataclass
class ExtractedDocument:
    text: str
    page_count: int = 0
    word_count: int = 0
    # True when a PDF yielded almost no text, which usually means it is a scan.
    looks_like_scan: bool = False
    warnings: list[str] = field(default_factory=list)


def normalise(text: str) -> str:
    """Clean text while preserving line structure.

    The old `clean_text` collapsed every run of whitespace including newlines,
    which destroyed the line breaks that section and name detection rely on.
    """
    text = unicodedata.normalize("NFKC", text)
    text = text.replace("•", " ")  # bullets
    text = re.sub(r"[ \t\r\f\v]+", " ", text)  # horizontal whitespace only
    text = re.sub(r"\n{3,}", "\n\n", text)  # cap blank runs
    text = "\n".join(line.strip() for line in text.split("\n"))
    return text.strip()


def _extract_pdf(path: Path) -> ExtractedDocument:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    if reader.is_encrypted:
        try:
            reader.decrypt("")
        except Exception as exc:  # noqa: BLE001 - surfaced to the user as a warning
            raise ExtractionError(f"PDF is password protected: {exc}") from exc

    pages: list[str] = []
    warnings: list[str] = []
    for index, page in enumerate(reader.pages):
        try:
            pages.append(page.extract_text() or "")
        except Exception as exc:  # noqa: BLE001 - one bad page should not kill the file
            warnings.append(f"Page {index + 1} could not be read ({exc}).")

    text = normalise("\n".join(pages))
    words = len(text.split())
    # A text-bearing resume page yields well over 40 words; below that it is a scan.
    looks_like_scan = len(reader.pages) > 0 and words < 40 * len(reader.pages)
    if looks_like_scan:
        warnings.append(
            "Very little selectable text was found — this looks like a scanned "
            "image. Export the resume as a text-based PDF so it can be read."
        )
    return ExtractedDocument(
        text=text,
        page_count=len(reader.pages),
        word_count=words,
        looks_like_scan=looks_like_scan,
        warnings=warnings,
    )


def _extract_docx(path: Path) -> ExtractedDocument:
    from docx import Document

    document = Document(str(path))
    blocks = [p.text for p in document.paragraphs]

    # Tables hold real content in many resume templates and were skipped before.
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                blocks.append(" | ".join(cells))

    text = normalise("\n".join(blocks))
    return ExtractedDocument(text=text, page_count=0, word_count=len(text.split()))


def _extract_txt(path: Path) -> ExtractedDocument:
    raw = path.read_text(encoding="utf-8", errors="replace")
    text = normalise(raw)
    return ExtractedDocument(text=text, word_count=len(text.split()))


def extract_text(path: str | Path) -> ExtractedDocument:
    """Extract text from a resume file, raising ExtractionError on failure."""
    path = Path(path)
    if not path.is_file():
        raise ExtractionError(f"File not found: {path.name}")

    suffix = path.suffix.lower()
    if suffix in UNSUPPORTED_LEGACY:
        raise ExtractionError(
            f"{suffix} is a legacy format that cannot be read directly. "
            "Please re-save the resume as PDF or DOCX."
        )
    if suffix not in SUPPORTED_SUFFIXES:
        raise ExtractionError(f"Unsupported file type: {suffix or 'unknown'}")

    try:
        if suffix == ".pdf":
            document = _extract_pdf(path)
        elif suffix == ".docx":
            document = _extract_docx(path)
        else:
            document = _extract_txt(path)
    except ExtractionError:
        raise
    except Exception as exc:  # noqa: BLE001 - normalise any parser failure
        raise ExtractionError(f"Could not read {path.name}: {exc}") from exc

    if not document.text.strip():
        raise ExtractionError(
            f"No text could be extracted from {path.name}. "
            "If this is a scanned document it needs OCR first."
        )
    return document
